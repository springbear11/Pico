#!/usr/bin/env python3
"""Generate PicoATE Architecture Design Document as a Word (.docx) file."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page Setup ──────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# ── Style Definitions ───────────────────────────────────────────

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.4

# Set East Asian font
rPr = style.element.get_or_add_rPr()
rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="微软雅黑"/>')
rPr.append(rFonts)

for i in range(1, 5):
    h_style = doc.styles[f'Heading {i}']
    h_font = h_style.font
    h_font.name = '微软雅黑'
    h_font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    h_rPr = h_style.element.get_or_add_rPr()
    h_rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="微软雅黑"/>')
    h_rPr.append(h_rFonts)
    if i == 1:
        h_font.size = Pt(22)
        h_style.paragraph_format.space_before = Pt(24)
        h_style.paragraph_format.space_after = Pt(16)
    elif i == 2:
        h_font.size = Pt(16)
        h_style.paragraph_format.space_before = Pt(20)
        h_style.paragraph_format.space_after = Pt(10)
    elif i == 3:
        h_font.size = Pt(13)
        h_style.paragraph_format.space_before = Pt(16)
        h_style.paragraph_format.space_after = Pt(8)
    elif i == 4:
        h_font.size = Pt(11)
        h_style.paragraph_format.space_before = Pt(12)
        h_style.paragraph_format.space_after = Pt(6)

# Code style
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_font = code_style.font
code_font.name = 'Consolas'
code_font.size = Pt(9)
code_font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
code_style.paragraph_format.space_before = Pt(2)
code_style.paragraph_format.space_after = Pt(2)
code_style.paragraph_format.line_spacing = 1.2
code_style.paragraph_format.left_indent = Cm(0.8)

# ── Helper Functions ────────────────────────────────────────────

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p

def add_code(text):
    lines = text.strip().split('\n')
    for line in lines:
        doc.add_paragraph(line, style='CodeBlock')
    return doc.paragraphs[-1]

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9.5)
    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacer
    return table

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))
    return p

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(f'[注] {text}')
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return p

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x88)
    return p

def page_break():
    doc.add_page_break()

# ── Cover Page ──────────────────────────────────────────────────

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('PicoATE 任务引擎')
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('架构设计文档')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x66)

doc.add_paragraph()
doc.add_paragraph()

info_lines = [
    ('版本', '1.0'),
    ('日期', '2026-06-26'),
    ('性质', '架构评审文档'),
    ('范围', 'PicoATE 源码库（截至 2026-06-26）'),
]
for label, value in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}：{value}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

page_break()

# ── Table of Contents (placeholder) ─────────────────────────────

add_heading('目录', 1)
add_para('（请在 Word 中右键此处 → 更新域 以生成自动目录）', italic=True, size=9)
# Insert TOC field
p = doc.add_paragraph()
run = p.add_run()
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run._r.append(fldChar1)
run2 = p.add_run()
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
run2._r.append(instrText)
run3 = p.add_run()
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
run3._r.append(fldChar2)
run4 = p.add_run('[ 更新此域以生成目录 ]')
run5 = p.add_run()
fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run5._r.append(fldChar3)

page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 1
# ═══════════════════════════════════════════════════════════════

add_heading('一、系统定位', 1)

add_heading('1.1 引擎身份', 2)
add_para('PicoATE（Pico Automated Test Equipment）是一个生产测试系统的任务调度与执行引擎，作为整个 ATE 软件栈中的核心中间层。它负责将声明式 JSON 测试序列编译为可执行计划，管理多 UUT 并发调度、资源仲裁、批次同步、错误恢复和设备连接生命周期，同时保持与上层 UI 和下层业务模块的解耦。')
add_para('源码依据（README.md）："PicoATE is the first implementation slice of the ATE runtime described in architecture_v3_1.md."')

add_heading('1.2 在整体系统中的位置', 2)
add_code("""┌──────────────────────────────────────────────┐
│            MES / 工厂系统（未来接入）           │
├──────────────────────────────────────────────┤
│         UI 层（Qt Widgets，规划中）            │
│     消费 ExecutionReport 只读 DTO             │
├──────────────────────────────────────────────┤
│          ★ PicoATE 任务引擎（本工程）★         │
│  ┌──────────┐ ┌──────────┐ ┌──────────────┐ │
│  │JSON编译器│ │ 调度引擎  │ │设备连接管理器 │ │
│  │+Plan     │ │Scheduler │ │DeviceSession  │ │
│  │Builder   │ │          │ │Manager        │ │
│  └──────────┘ └──────────┘ └──────────────┘ │
│  ┌──────────────────────────────────────────┐│
│  │      ModuleRegistry / NodeRunner         ││
│  └──────────────────────────────────────────┘│
├──────────────────────────────────────────────┤
│             Transport 抽象层                  │
│  QProcess │ Persistent │ DllBridge │ Native  │
├──────────────────────────────────────────────┤
│        业务模块（DLL / Python / EXE）          │
├──────────────────────────────────────────────┤
│       设备层（DMM / CAN / PSU / Fixture）      │
└──────────────────────────────────────────────┘""")

add_heading('1.3 解决的问题', 2)
add_table(
    ['问题域', '引擎承担的职责', '关键源码'],
    [
        ['流程编排', 'JSON 描述 Setup → Main → Cleanup 序列', 'SequenceCompiler.h:30'],
        ['资源仲裁', '多 UUT 对 DMM1/CAN1 的互斥申请与排队', 'ResourceManager.h:62'],
        ['批次同步', '多 UUT 到 Barrier 点后等待、到齐释放', 'BarrierController.h:102'],
        ['错误恢复', '失败→重试、失败→清理、失败→跳过', 'ErrorPolicyEngine.h:16'],
        ['循环控制', 'For 循环由调度层控制，Plan 保持不可变', 'LoopController.h:14'],
        ['设备连接', '逻辑设备 ID ↔ 真实地址映射、连接复用', 'DeviceSessionManager.h:77'],
        ['模块隔离', 'DLL/Python 崩溃不影响引擎进程', 'nativehost/Main.cpp'],
    ],
    [3.5, 9, 6.5]
)

add_heading('1.4 输入与输出', 2)
add_para('输入：', bold=True)
add_table(
    ['输入', '格式', '提供方'],
    [
        ['Sequence JSON', 'JSON 文件，描述测试步骤、策略、输入', '测试工程师/项目团队'],
        ['Station Config JSON', 'JSON 文件，描述工站设备地址、driver', '工站配置'],
        ['UUT 数量', 'CLI 参数 --uuts N', '操作员/上位机'],
        ['业务模块', 'DLL（C ABI）、Python 脚本、独立 EXE', '项目团队'],
    ],
    [4, 8.5, 6.5]
)

add_para('输出：', bold=True)
add_table(
    ['输出', '格式', '消费者'],
    [
        ['ExecutionReport', '结构化 DTO（UUT/Step/Attempt/Measurement）', 'UI / CLI / MES'],
        ['ExecutionSessionResult', '含 completed、hasError、nodeResults', 'CLI / 调用方'],
        ['SessionSnapshot', '可序列化的资源/Barrier/UUT 状态', '未来持久化/恢复'],
        ['CLI 文本输出', 'stdout 打印 step/attempt/measurement', '调试/日志'],
    ],
    [5, 8.5, 5.5]
)

add_heading('1.5 上下游依赖', 2)
add_table(
    ['外部系统', '当前状态', '交互方式'],
    [
        ['UI 层', '无，已规划', '通过 ExecutionReport DTO（只读）消费'],
        ['设备层 (DMM/CAN/PSU)', 'Fake 已验证，真实硬件未接', 'IModuleTransport → PersistentQProcessTransport'],
        ['MES', '无', '当前代码无法确认'],
        ['数据库', '无', '仅 ISessionPersistence 接口占位'],
        ['Python 脚本', '已验证', 'QProcessTransport + moduleBindings'],
        ['C/C++ DLL', '已验证', 'PicoATE.NativeHost.exe + PicoATE_Execute C ABI'],
    ],
    [4, 5.5, 9.5]
)

add_heading('1.6 核心设计约束', 2)
add_para('来自 project_vision.md 的架构护栏：')
add_table(
    ['约束', '含义'],
    [
        ['协议逻辑不进调度器', 'CAN/LIN/DMM 命令解析属于模块或 Transport Host'],
        ['UI 只消费只读 DTO', '使用 ExecutionReport，不读 NodeActivation 内部状态'],
        ['调度器拥有编排权', 'Retry/Cleanup/Loop/Barrier/Resource 集中管理'],
        ['模块返回中性结果', '返回 ModuleResult，不返回 NodeResult 或 UI 数据'],
        ['配置声明式', '流程/策略/资源/输入用 JSON 描述，不编译进代码'],
        ['Transport 可替换', 'DLL/Python/QProcess/gRPC 都放 IModuleTransport 后面'],
    ],
    [5, 14]
)

page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 2
# ═══════════════════════════════════════════════════════════════

add_heading('二、项目目录结构分析', 1)

add_heading('2.1 顶层目录', 2)
add_code("""PicoATE/
├── CMakeLists.txt           # C++20, Qt6 Core+Test
├── CMakePresets.json        # VS2022 + Qt6.9.1 x64 预设
├── README.md                # 构建/运行说明
├── test_output.txt          # 最近测试输出记录
├── docs/                    # 15 份设计/规范/总结文档
├── examples/                # 16 个示例（JSON序列、模块、工站配置）
├── src/                     # 全部源代码
│   ├── core/                # ★ 核心引擎库（静态库，27头+27实现）
│   ├── cli/                 # 命令行入口
│   ├── mockhost/            # 测试用外部 echo 进程
│   ├── nativehost/          # DLL 隔离加载器进程
│   ├── fakeinstrumenthost/  # 假仪器长驻进程
│   ├── testdllmodule/       # C ABI 测试 DLL
│   └── canexamplemodule/    # 模拟 CAN 解码 DLL
├── tests/                   # 单元测试（3408行，34用例）
└── out/                     # 构建输出""")

add_heading('2.2 核心模块职责', 2)

add_heading('2.2.1 数据模型层（src/core/）', 3)
add_table(
    ['头文件', '核心类型', '职责'],
    [
        ['ExecutionPlan.h', 'ExecNode, ExecEdge, ExecutionPlan, CleanupRegion, LoopRegion, ForLoopSpec', '不可变执行计划'],
        ['RuntimeTypes.h', 'UutExecution, NodeActivation, NodeAttempt, ExecutionState, ActivationState', '运行时可变状态'],
        ['SequenceDef.h', 'SequenceDef, StepDef, StepGroupDef, ModuleBindingDef', '编辑期模型'],
        ['MeasurementTypes.h', 'MeasurementResult, MeasurementStatus', '结构化测量值 DTO'],
        ['ExecutionReport.h', 'ExecutionReport, UutReport, StepReport, AttemptReport', '只读运行报告 DTO'],
        ['SessionSnapshot.h', 'ExecutionSessionSnapshot, ISessionPersistence', '会话快照/持久化接口占位'],
    ],
    [4.5, 8, 6.5]
)

add_heading('2.2.2 编译与构建层', 3)
add_table(
    ['头文件', '核心类型', '职责'],
    [
        ['SequenceCompiler.h', 'SequenceCompiler, CompileResult, CompileError, CompileWarning', 'JSON → SequenceDef → ExecutionPlan'],
        ['PlanBuilder.h', 'PlanBuilder, PlanBuildResult', 'SequenceDef → ExecutionPlan 图构建'],
        ['PlanCache.h', 'PlanCache', '基于 shared_ptr 的编译缓存'],
    ],
    [4.5, 8, 6.5]
)

add_heading('2.2.3 调度引擎层', 3)
add_table(
    ['头文件', '核心类型', '职责'],
    [
        ['ExecutionSession.h', 'ExecutionSession, ExecutionSessionResult', '多 UUT 执行会话主入口'],
        ['ExecutionGraphScheduler.h', 'ExecutionGraphScheduler, SchedulerResult', '节点调度主循环'],
        ['ResourceManager.h', 'ResourceManager, ResourceLease', '资源仲裁（4种模式）'],
        ['BarrierController.h', 'BarrierController, BarrierReleaseDecision', '批次同步（6×4×5策略）'],
        ['LoopController.h', 'LoopController', 'For 循环游标控制'],
        ['ErrorPolicyEngine.h', 'ErrorPolicyEngine, ErrorDecision', '错误策略决策（5种动作）'],
        ['NodeRunner.h', 'NodeRunner, INodeHandler, ActionNodeHandler', '节点分发执行器'],
    ],
    [4.5, 8, 6.5]
)

add_heading('2.2.4 模块运行时与 Transport 层', 3)
add_table(
    ['头文件', '核心类型', '职责'],
    [
        ['ModuleRuntime.h', 'IModule, IModuleTransport, ModuleRegistry, TransportModuleAdapter', '模块执行与传输抽象'],
        ['ModuleTransportJson.h', '序列化函数', 'Transport 请求/响应 ↔ JSON'],
        ['ModuleBindingRegistrar.h', 'registerConfiguredModules()', 'JSON bindings → 运行时注册'],
        ['QProcessTransport.h', 'QProcessTransport', '短生命周期子进程 transport'],
        ['PersistentQProcessTransport.h', 'PersistentQProcessTransport', '长驻子进程 transport'],
        ['DllBridgeInvoker.h', 'DllBridgeInvoker', '进程内 QLibrary DLL 调用'],
        ['NativeHostManifest.h', 'NativeHostManifest', 'DLL 加载清单 JSON 解析'],
    ],
    [4.5, 8, 6.5]
)

add_heading('2.2.5 设备管理与其他', 3)
add_table(
    ['头文件', '核心类型', '职责'],
    [
        ['DeviceSessionManager.h', 'DeviceSessionManager, IDeviceSession, IDeviceSessionFactory, DeviceSessionConfig', '逻辑设备注册/连接复用'],
        ['DeviceTransportSession.h', 'TransportDeviceSession, TransportDeviceSessionFactory', 'Transport 映射为设备 session'],
        ['StationConfig.h', 'StationConfig, parseStationConfigJson()', '工站配置 JSON 解析'],
        ['StationRuntime.h', 'StationRuntime', '工站运行上下文'],
        ['InstrumentAdapterModules.h', 'ExampleDmmAdapterModule, ExampleCanAdapterModule', 'DMM/CAN 适配器 Spike'],
        ['VariableResolver.h', 'VariableResolver', '配置期变量替换'],
        ['RuntimeVariableResolver.h', 'RuntimeVariableResolver', '运行期变量替换'],
    ],
    [4.5, 8, 6.5]
)

add_heading('2.3 各可执行程序', 2)
add_table(
    ['程序', '文件', '用途'],
    [
        ['PicoATE.Cli', 'src/cli/Main.cpp (517行)', '命令行入口，组装完整 pipeline'],
        ['PicoATE.MockHost', 'src/mockhost/Main.cpp (72行)', '测试用 stdin/stdout echo 进程'],
        ['PicoATE.NativeHost', 'src/nativehost/Main.cpp (207行)', 'DLL 隔离加载器，支持 --manifest'],
        ['PicoATE.FakeInstrumentHost', 'src/fakeinstrumenthost/Main.cpp (320行)', '假仪器长驻进程，验证跨 step 状态保持'],
        ['PicoATE.TestDllModule', 'src/testdllmodule/TestDllModule.cpp (80行)', 'C ABI 测试 DLL'],
        ['PicoATE.CanExampleModule', 'src/canexamplemodule/CanExampleModule.cpp (247行)', '模拟 CAN 解码 DLL'],
    ],
    [5, 6, 8]
)

add_heading('2.4 模块依赖关系', 2)
add_code("""                   PicoATECli
                       │
      ┌────────────────┼────────────────┐
      ▼                ▼                ▼
 PicoATECore    PicoATEMockHost   PicoATENativeHost
  (静态库)       (独立进程)         (独立进程)
      │
      └── 内部层次（自底向上）：

 Transport实现 → IModuleTransport(抽象)
                 → IModule(抽象)
                   → NodeRunner → Scheduler → Session
                                             → StationRuntime
                                               → CLI

 上层依赖抽象，具体实现通过 ModuleRegistry / DeviceSessionManager 注册注入。""")

page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 3
# ═══════════════════════════════════════════════════════════════

add_heading('三、核心架构设计', 1)

add_heading('3.1 架构风格识别', 2)
add_para('PicoATE 采用了"分层 + 管道 + 策略 + 注册"的混合架构风格：')
add_table(
    ['架构风格', '体现位置', '说明'],
    [
        ['分层架构', '编译层→会话层→调度层→模块层→Transport层', '依赖方向单向，上层不依赖下层具体实现'],
        ['管道 (Pipeline)', 'JSON→SequenceDef→ExecutionPlan→ExecutionSession→ExecutionReport', 'SequenceCompiler → PlanBuilder → ExecutionSession 链式处理'],
        ['策略模式', 'Barrier 6种到达策略、4种释放策略、Error 5种动作', '可从 JSON 配置切换，不需改调度器代码'],
        ['注册/插件', 'IModule 通过 ModuleRegistry 注册，IDeviceSessionFactory 通过 DeviceSessionManager 注册', '支持运行时动态注册，编译期不需要知道所有实现'],
        ['状态机', 'ExecutionState(9)、ActivationState(11)、FrameState(4)、AttemptState(4)', 'RuntimeTypes.h 定义四层状态机'],
    ],
    [3, 6.5, 9.5]
)

add_heading('3.2 数据流', 2)
add_code("""JSON 文件
  → SequenceCompiler.compileJson()
    → SequenceDef（编辑期模型）
      → PlanBuilder.build()
        → ExecutionPlan（不可变执行图）
          → ExecutionSession
            → 多 UUT 并行推进
              → ExecutionGraphScheduler.pumpOnce()
                → NodeRunner.run()
                  → IModule.execute()
                    → ModuleResult → NodeResult → NodeActivation
                      → ExecutionReport（只读 DTO）""")

add_para('关键设计决策：ExecutionPlan 构建后不可变。所有运行时状态（UutExecution, NodeActivation, NodeAttempt）存储在 ExecutionSession 中，不修改 Plan。这保证了同一个 Plan 可被多个 UUT 安全共享，PlanCache 可以安全缓存。源码证据：PlanCache.h:17 使用 shared_ptr<const ExecutionPlan>。')

add_heading('3.3 控制流', 2)
add_code("""ExecutionSession::run()
  └── 循环直到所有 UUT 完成：
        └── 对每个 UUT：
              └── ExecutionGraphScheduler::pumpOnce(uut)
                    ├── findReadyNodes()          // 拓扑找就绪节点
                    │     ├── 检查依赖是否满足
                    │     ├── 检查 Loop body 是否可运行
                    │     └── 跳过已终止节点
                    ├── executeNode()
                    │     ├── Loop 节点 → executeLoopNode()
                    │     │     └── LoopController::advance()
                    │     ├── Barrier 节点 → executeBarrierNode()
                    │     │     └── BarrierController::memberArrived()
                    │     └── Action/Cleanup 节点：
                    │           ├── ResourceManager::tryAcquire()
                    │           ├── NodeRunner::run()
                    │           │     └── RuntimeVariableResolver
                    │           │         └── IModule::execute()
                    │           ├── ErrorPolicyEngine::decide()
                    │           ├── activateCleanup()
                    │           ├── skipPendingNonAlwaysRun()
                    │           └── handleNodeFailureForBarriers()
                    └── applyBarrierReleases()""")

add_heading('3.4 调度模型', 2)
add_para('PicoATE 采用同步协作式调度：所有执行在主线程中通过循环泵送（pumpOnce）完成。每个 UUT 每轮尝试推进一个就绪节点。')
add_table(
    ['特征', '说明', '源码证据'],
    [
        ['单线程', '没有后台线程池执行测试步骤', 'ExecutionSession.cpp:260-281 的 while 循环在主线程'],
        ['协作式', '每个节点执行完后返回，不长时间阻塞主循环', 'NodeRunner.cpp 中各 Handler 的同步返回'],
        ['跨 UUT 轮转', 'Round-robin 方式轮流推进各个 UUT', 'ExecutionSession.cpp:264 for(auto& uut : m_uuts)'],
    ],
    [2.5, 10.5, 6]
)
add_note('当前代码无法确认：是否有计划引入多线程并发执行模型。从 project_vision.md 中也未看到相关规划。')

add_heading('3.5 并发控制', 2)
add_table(
    ['机制', '实现方式', '代码位置'],
    [
        ['资源互斥', 'ResourceManager 的 tryAcquire/release，支持 Exclusive/SharedRead/SharedWrite/Counted', 'ResourceManager.h:62'],
        ['等待队列', '资源申请失败时加入 waiter 队列，资源释放不自动唤醒', 'ResourceManager.h:76-77'],
        ['Barrier 联合', '多 UUT 到 Barrier 点后等待，满足条件后统一释放', 'BarrierController.h:102'],
        ['停止保护', 'ExecutionSession 的 stopRequested 标志 + skipPendingNonAlwaysRun', 'ExecutionSession.cpp:370-382'],
    ],
    [3, 11, 5]
)
add_note('当前代码无法确认：是否有多线程锁（QMutex）、条件变量（QWaitCondition）、原子操作。当前引擎是单线程设计，没有多线程并发。')

page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 4
# ═══════════════════════════════════════════════════════════════

add_heading('四、任务生命周期分析', 1)

add_heading('4.1 Session 状态机', 2)
add_para('ExecutionState（9 个状态，RuntimeTypes.h:12-23）：Idle → Running → (Stopping → CleaningUp) → Completed/CompletedWithError/Aborted')
add_code("""  Idle ──→ Running ──→ Completed
    │          │
    │          ├──→ Stopping ──→ CleaningUp ──→ Completed / CompletedWithError
    │          │
    │          └──→ Aborted
    │
    └──→ (Paused — 枚举已定义但逻辑未实现)""")

add_heading('4.2 节点状态机 (ActivationState)', 2)
add_para('ActivationState（11 个状态，RuntimeTypes.h:45-58）：')
add_code("""Created → WaitingForDependency → Ready
                           ↓
                     WaitingForResource ──→ Running
                           ↓                  ↓
                     (重试) ←── Failed/Error/Timeout
                                              ↓
                  WaitingAtBarrier → Passed/Cancelled/Skipped""")

add_heading('4.3 完整生命周期 9 阶段', 2)
add_code("""阶段 1：编译
  JSON → SequenceCompiler.compileJson() → SequenceDef
  含类型校验、未知字段警告

阶段 2：构建
  SequenceDef → PlanBuilder.build() → ExecutionPlan（不可变）
  生成节点、边、CleanupRegion、LoopRegion

阶段 3：会话初始化
  ExecutionSession(plan)
  创建 Scheduler、ResourceManager、BarrierController、LoopController

阶段 4：UUT 注册
  addUut("UUT-1"), addUut("UUT-2") → 创建 UutExecution 对象

阶段 5：模块注册
  registerConfiguredModules()
  解析 moduleBindings → 创建 Transport → 注册 Adapter

阶段 6：Station 配置（可选）
  StationRuntime.loadStationConfigFile()
  配置 DeviceSessionManager（DMM1/CAN1/PSU1 等）

阶段 7：执行循环
   while (progressed):
     for each UUT: pumpOnce(uut)
       → findReadyNodes() → executeNode()
          → 资源获取 → 模块执行（含变量替换）
          → 错误决策 → 重试/清理/跳过 → 释放资源
       → applyBarrierReleases()
     break if allUutsComplete()

阶段 8：报告生成
  session.report() → ExecutionReport

阶段 9：结果消费
  CLI 打印 / UI 渲染 / 未来 MES 上传""")

add_heading('4.4 设备连接生命周期', 2)
add_code("""工站启动 → StationRuntime.loadStationConfigFile()
  → 配置 DeviceSessionManager

运行开始 → 业务模块 ConnectDMM
  → runtimeServices->openDeviceSession("DMM1")
    → 创建 session, connect()

测试步骤 → ConfigureDMM / ReadDMM（多次）
  → 复用已连接 session

Cleanup → DisconnectDMM
  → session->disconnect()，session 对象保留

后续运行 → 再次 ConnectDMM
  → 复用已存在的 session 对象，只重新 connect()""")
add_para('源码证据：DeviceSessionManager.cpp 实现了连接复用语义。测试验证见 CoreTests.cpp 的 deviceSessionManagerReusesConnectedSession 用例。')

page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 5
# ═══════════════════════════════════════════════════════════════

add_heading('五、三层解耦架构分析', 1)

add_heading('5.1 架构分层定义', 2)
add_para('源码 project_vision.md:39-55 明确定义了三层架构：')
add_code("""┌──────────────────────────────────────────────┐
│               UI 层（规划中）                   │
│  编辑配置 / 启动停止执行 / 显示报告              │
│  消费 ViewModel / ExecutionReport DTO           │
│  不得包含产品测试逻辑                            │
├──────────────────────────────────────────────┤
│          ★ 任务引擎层（Scheduler Layer）★        │
│  JSON→SequenceDef→ExecutionPlan→Session         │
│  拥有 Retry, Cleanup, Barrier, Resource 策略    │
│  通过 IModule / IModuleTransport 调用业务        │
│  不得包含 CAN/DLL/仪器/产品特定逻辑              │
├──────────────────────────────────────────────┤
│            业务测试逻辑层                        │
│  实现 IModule 或外部模块 Host                    │
│  拥有协议解析、仪器控制、测量、项目逻辑           │
│  返回 ModuleResult / ModuleTransportResponse    │
└──────────────────────────────────────────────┘""")

add_heading('5.2 各层职责分析', 2)

add_para('UI 层：', bold=True)
add_para('当前状态：未实现。只有 ExecutionReport 作为 UI 消费的 DTO 已就绪，以及 project_vision.md 中的设计规划。CLI 是当前唯一的"用户界面"。')
add_note('当前代码无法确认：UI 具体框架、UI 如何获取 ExecutionReport、UI 如何发送控制命令。')

add_para('任务引擎层（完成度约 70%-75%）：', bold=True)
add_bullet('JSON 编译与校验：SequenceCompiler + PlanBuilder')
add_bullet('不可变执行图：ExecutionPlan')
add_bullet('多 UUT 调度：ExecutionSession + ExecutionGraphScheduler')
add_bullet('资源仲裁：ResourceManager（4 种模式）')
add_bullet('批次同步：BarrierController（6×4×5 策略组合）')
add_bullet('循环控制：LoopController（For 循环）')
add_bullet('错误恢复：ErrorPolicyEngine（5 种错误动作）')
add_bullet('引擎层依赖的都是抽象接口：IModule、IModuleTransport、IDeviceSession、IDeviceSessionFactory')

add_para('业务逻辑层（完成度约 82%-87%）：', bold=True)
add_table(
    ['接入方式', '隔离程度', '已验证'],
    [
        ['内置模块 (IModule 直接注册)', '进程内', 'mock.action, mock.measurement, example.dmm, example.can'],
        ['QProcessTransport（短进程）', '进程外', 'PicoATE.MockHost, Python echo_module.py'],
        ['PersistentQProcessTransport（长驻进程）', '进程外', 'PicoATE.FakeInstrumentHost'],
        ['NativeHost + DLL（C ABI）', '进程外', 'TestDllModule, CanExampleModule'],
        ['DllBridgeInvoker（直接 QLibrary）', '进程内（验证用）', 'TestDllModule'],
    ],
    [4, 6, 9]
)

add_heading('5.3 依赖方向检查', 2)
add_table(
    ['检查项', '状态', '说明'],
    [
        ['UI → 只依赖 ExecutionReport DTO', '✓ 通过', '规划中，当前无 UI'],
        ['UI → 不依赖 IModule/业务实现', '✓ 通过', '规划中'],
        ['引擎 → 只依赖 IModule/IModuleTransport/IDeviceSession 抽象', '✓ 通过', '所有调用经抽象接口'],
        ['引擎 → 不依赖具体 CAN/DMM/仪器实现', '✓ 通过', 'example.* 仅用于 Spike'],
        ['业务 → 不依赖 UI', '✓ 通过', '业务返回 ModuleResult，不知 UI 存在'],
        ['业务 → 不依赖 Scheduler 内部状态', '✓ 通过', '接收 ModuleExecutionContext，返回 ModuleResult'],
    ],
    [7.5, 2.5, 9]
)

add_para('关键证据：IModule 接口只接受 ModuleExecutionContext 和 ModuleFunction，返回 ModuleResult（ModuleRuntime.h:107-113）。引擎层负责将 ModuleOutcome 转换为 NodeOutcome（ModuleRuntime.h:157）：toNodeOutcome() 函数完成了这层隔离。', bold=False)

page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 6
# ═══════════════════════════════════════════════════════════════

add_heading('六、任务引擎提供给 UI 的接口', 1)

add_heading('6.1 当前状态', 2)
add_para('UI 层尚未实现。当前只有 CLI 作为调用入口。以下是当前代码中引擎对外暴露的实际 API。')

add_heading('6.2 当前 API 清单', 2)
add_table(
    ['接口类别', '接口', '输入', '输出', '异步', '作用', '源码'],
    [
        ['编译', 'SequenceCompiler::compileJson()', 'QJsonObject', 'CompileResult', '否', 'JSON → SequenceDef + Plan', 'SequenceCompiler.h:31'],
        ['构建', 'PlanBuilder::build()', 'SequenceDef', 'PlanBuildResult', '否', 'SequenceDef → Plan', 'PlanBuilder.h:21'],
        ['会话创建', 'ExecutionSession()', 'ExecutionPlan', '-', '否', '创建执行会话', 'ExecutionSession.h:20'],
        ['UUT 注册', 'addUut()', 'UutId', 'UutExecution&', '否', '添加被测试单元', 'ExecutionSession.h:22'],
        ['模块注册', 'registerModule()', 'shared_ptr<IModule>', 'bool', '否', '注册业务模块', 'ExecutionSession.h:27'],
        ['运行', 'run()', '-', 'ExecutionSessionResult', '否', '启动执行', 'ExecutionSession.h:32'],
        ['停止', 'requestStop()', 'StopMode', 'void', '否', '请求停止', 'ExecutionSession.h:29'],
        ['状态查询', 'state()', '-', 'ExecutionState', '否', '查询会话状态', 'ExecutionSession.h:30'],
        ['报告', 'report()', '-', 'ExecutionReport', '否', '生成只读报告', 'ExecutionSession.h:33'],
        ['快照', 'snapshot()', '-', 'ExecutionSessionSnapshot', '否', '生成序列化快照', 'ExecutionSession.h:34'],
        ['设备管理', 'devices()', '-', 'DeviceSessionManager&', '否', '设备管理器引用', 'ExecutionSession.h:24'],
        ['UUT 查询', 'uuts()', '-', 'QVector<UutExecution>&', '否', 'UUT 执行数据', 'ExecutionSession.h:23'],
        ['工站加载', 'loadStationConfigFile()', 'filePath, options', 'StationRuntimeResult', '否', '加载工站配置', 'StationRuntime.h:16'],
        ['模块注册', 'registerConfiguredModules()', 'session, sequence, options', 'RegistrationResult', '否', '从序列注册模块', 'ModuleBindingRegistrar.h:29'],
    ],
    [2, 4, 3, 5, 1, 3, 2.5]
)

add_heading('6.3 存在的问题', 2)
add_bullet('没有统一的 UI 门面接口（Facade）：CLI 直接操作多个底层对象（SequenceCompiler、PlanBuilder、ExecutionSession、StationRuntime、ModuleBindingRegistrar）')
add_bullet('当前代码无法确认：是否有事件/回调机制用于通知 UI 状态变化（当前是轮询模式）')
add_bullet('当前代码无法确认：Pause/Resume 的实现（ExecutionState 中有 Paused 枚举，但 run() 中无对应逻辑）')
add_bullet('当前代码无法确认：是否有进度回调（当前只能通过 state() 轮询）')
add_bullet('当前代码无法确认：是否有日志订阅接口')

page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 7
# ═══════════════════════════════════════════════════════════════

add_heading('七、任务引擎提供给业务层的接口', 1)

add_heading('7.1 IModule — 业务模块核心接口', 2)
add_code("""class IModule {
public:
    virtual ModuleId moduleId() const = 0;
    virtual ModuleResult execute(
        const ModuleFunction& functionName,
        const ModuleExecutionContext& context) = 0;
};  // 源码：ModuleRuntime.h:107-113""")

add_heading('7.2 IModuleTransport — 外部进程/DLL 传输接口', 2)
add_code("""class IModuleTransport {
public:
    virtual ModuleTransportStatus call(
        const ModuleTransportRequest& request,
        ModuleTransportResponse& response,
        int timeoutMs) = 0;
};  // 源码：ModuleRuntime.h:99-105""")

add_heading('7.3 IModuleRuntimeServices — 设备服务接口', 2)
add_code("""class IModuleRuntimeServices {
public:
    virtual DeviceSessionOpenResult openDeviceSession(const DeviceId&) = 0;
    virtual DeviceSessionError closeDeviceSession(const DeviceId&) = 0;
    virtual shared_ptr<IDeviceSession> deviceSession(const DeviceId&) const = 0;
    virtual ModuleResult invokeDevice(const DeviceId&,
        const ModuleFunction&, const QVariantMap&,
        const ModuleExecutionContext&) = 0;
};  // 源码：ModuleRuntime.h:49-60""")

add_heading('7.4 业务层接口清单汇总', 2)
add_table(
    ['接口', '实现方', '调用方', '输入', '输出', '作用'],
    [
        ['IModule::moduleId()', '业务模块', 'ModuleRegistry', '-', 'ModuleId', '模块标识'],
        ['IModule::execute()', '业务模块', 'ActionNodeHandler', 'functionName, context', 'ModuleResult', '执行测试'],
        ['IModuleTransport::call()', 'Transport实现', 'TransportModuleAdapter', 'request, timeoutMs', 'Status+Response', '跨进程调用'],
        ['IModuleRuntimeServices::openDeviceSession()', 'ModuleRuntimeServices', '业务模块', 'deviceId', 'OpenResult', '打开设备'],
        ['IModuleRuntimeServices::closeDeviceSession()', 'ModuleRuntimeServices', '业务模块', 'deviceId', 'SessionError', '关闭设备'],
        ['IModuleRuntimeServices::invokeDevice()', 'ModuleRuntimeServices', '业务模块', 'deviceId,function,inputs,context', 'ModuleResult', '设备命令'],
        ['PicoATE_Execute() (C ABI)', 'DLL模块', 'NativeHost/DllBridge', 'requestJson,buffer,size', 'int(0=成功)', 'DLL调用'],
    ],
    [4, 3, 3, 4, 3, 3]
)

add_heading('7.5 接口边界检查', 2)
add_table(
    ['检查项', '状态'],
    [
        ['业务模块不接触 NodeResult/NodeActivation/ExecutionPlan', '✓ 通过'],
        ['业务模块通过 ModuleExecutionContext 获取上下文', '✓ 通过'],
        ['业务模块通过 IModuleRuntimeServices 访问设备', '✓ 通过'],
        ['业务模块不直接操作 DeviceSessionManager', '✓ 通过'],
        ['Transport 层完全对业务模块透明', '✓ 通过'],
    ],
    [10, 6]
)

page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 8
# ═══════════════════════════════════════════════════════════════

add_heading('八、三层交互方式分析', 1)

add_heading('8.1 通信方式总结', 2)
add_table(
    ['层次间', '通信方式', '实现'],
    [
        ['UI → 引擎', '直接 API 调用（当前）/ 未来 Event/Command', 'ExecutionSession 公开方法'],
        ['引擎 → UI', '轮询 report()（当前）/ 未来 Callback/Event', '当前无推送机制'],
        ['引擎 → 业务', '虚函数调用（IModule::execute）', 'ActionNodeHandler.run()'],
        ['业务 → 引擎', '函数返回 ModuleResult', 'execute() 的返回值'],
        ['引擎 → Transport', '虚函数调用（IModuleTransport::call）', 'TransportModuleAdapter'],
        ['Transport → 外部进程', 'stdin/stdout JSON 行协议', 'QProcessTransport'],
        ['引擎 → 设备层', '虚函数调用（IDeviceSession）', 'ModuleRuntimeServices'],
        ['业务 → 设备层', '通过 IModuleRuntimeServices', 'context.runtimeServices->invokeDevice()'],
    ],
    [3.5, 7, 8]
)

add_heading('8.2 正常执行流程', 2)
add_code("""UI/CLI                  ExecutionSession       Scheduler              NodeRunner           IModule(业务)
  │                          │                     │                      │                    │
  │──run()──────────────────→│                     │                      │                    │
  │                          │──pumpOnce(uut)─────→│                      │                    │
  │                          │                     │──findReadyNodes()    │                    │
  │                          │                     │──run(node, context)──→│                    │
  │                          │                     │                      │──RuntimeVariable   │
  │                          │                     │                      │──module(moduleId)  │
  │                          │                     │                      │──execute(fn, ctx)──→│
  │                          │                     │                      │                    │──业务逻辑
  │                          │                     │                      │←──ModuleResult─────│
  │                          │                     │←──NodeResult─────────│                    │
  │                          │                     │──ErrorPolicy.decide()│                    │
  │                          │                     │  (Retry→重新run)     │                    │
  │                          │←──SchedulerStepResult│                    │                    │
  │                          │──applyBarrierReleases│                    │                    │
  │  (循环直到所有UUT完成)    │                     │                      │                    │
  │←──ExecutionSessionResult│                     │                      │                    │
  │──report()──────────────→│                     │                      │                    │
  │←──ExecutionReport───────│                     │                      │                    │""")

add_heading('8.3 停止传播流程', 2)
add_code("""UI/CLI                  ExecutionSession       Scheduler
  │                          │                     │
  │──requestStop(Graceful)──→│                     │
  │                          │ m_stopRequested=true│
  │                          │ m_state = Stopping  │
  │                          │ prepareStopIfRequested()
  │                          │──skipPendingNonAlwaysRun(uut)
  │                          │  (将所有非alwaysRun节点标记Skipped)
  │                          │──activateAllCleanup(uut)
  │                          │  (激活所有Cleanup节点)""")

add_heading('8.4 失败传播流程', 2)
add_code("""Scheduler                 ErrorPolicyEngine      BarrierController
  │                          │                     │
  │ 模块返回 Failed          │                     │
  │──decide(node,result)────→│                     │
  │←──ErrorDecision─────────│                     │
  │  (action=RunCleanup)     │                     │
  │──activateCleanup()       │                     │
  │──handleNodeFailureForBarriers()
  │  (通知后续Barrier有成员失败)
  │──memberFailedBeforeArrival(uutId, barrierId)──→│
  │←──BarrierReleaseDecision──────────────────────│
  │──skipPendingNonAlwaysRun()
  │  (跳过该UUT所有后续非alwaysRun节点)""")

add_heading('8.5 暂停/恢复', 2)
add_note('当前代码无法确认：ExecutionState 枚举中定义了 Paused，但在 ExecutionSession::run() 中没有对应的暂停/恢复逻辑。Pause/Resume 是规划中但未实现的能力。')

page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 9
# ═══════════════════════════════════════════════════════════════

add_heading('九、核心流程分析', 1)

add_heading('9.1 调度流程', 2)
add_code("""run() 开始
  ├── m_stopRequested?
  │     ├── 是 → prepareStopIfRequested（跳过非alwaysRun + 激活Cleanup）
  │     └── 否 → 继续
  ├── 遍历所有 UUT
  │     └── pumpOnce(uut)
  │           ├── findReadyNodes → 有就绪节点?
  │           │     ├── 否 → blocked=true, 下一轮重试
  │           │     └── 是 → 判断节点类型
  │           │           ├── Loop → executeLoopNode → LoopController::advance
  │           │           ├── Barrier → executeBarrierNode → memberArrived
  │           │           └── Action/Cleanup
  │           │                 ├── 有资源需求? → tryAcquire
  │           │                 │     ├── 失败 → WaitingForResource, 下一轮重试
  │           │                 │     └── 成功 → NodeRunner::run
  │           │                 └── 无资源 → NodeRunner::run
  │           │                       └── ErrorPolicyEngine::decide
  │           │                             ├── Retry → 重新 run
  │           │                             ├── RunCleanup → activateCleanup
  │           │                             └── StopUut/Abort → skipPendingNonAlwaysRun
  │           └── 释放资源 + applyBarrierReleases
  └── allUutsComplete? → 否→继续循环, 是→设置最终状态→返回""")

add_heading('9.2 数据流（编译 → 执行 → 报告）', 2)
add_code("""编译阶段:              执行阶段:                  报告阶段:
  JSON 文件                ExecutionPlan             ExecutionReport
    ↓                          ↓                         ↑
  SequenceCompiler      ExecutionSession             UutReport
    ↓                          ↓                         ↑
  SequenceDef           UutExecution×N             StepReport
    ↓                          ↓                         ↑
  PlanBuilder           NodeActivation×M          AttemptReport
    ↓                          ↓                     (含measurement
  ExecutionPlan          NodeAttempt×R               +loop)
  (不可变)              (含MeasurementResult)""")

add_heading('9.3 跨 UUT Barrier 时序', 2)
add_code("""UUT-1                UUT-2                Scheduler             BarrierController
  │                    │                    │                     │
  │                    │                    │──createBarrier()───→│
  │                    │                    │  ("batch-ready",    │
  │                    │                    │   {UUT-1,UUT-2})    │
  │                    │                    │                     │
  │ batch-ready就绪    │                    │                     │
  │──executeBarrierNode│                    │                     │
  │                    │                    │──memberArrived(1)──→│
  │                    │                    │←─not released──────│
  │ WaitingAtBarrier   │                    │                     │
  │                    │ batch-ready就绪    │                     │
  │                    │──executeBarrierNode│                    │
  │                    │                    │──memberArrived(2)──→│
  │                    │                    │←─released!─────────│
  │                    │                    │  {UUT-1,UUT-2}      │
  │                    │                    │                     │
  │←─Passed────────────│←─Passed────────────│                     │
  │ 继续after-barrier  │ 继续after-barrier  │                     │""")

add_heading('9.4 For 循环执行流程', 2)
add_code("""  LoopController: sampleIndex=0..2 step=1

  第0轮: advance() → sampleIndex=0 → executeNode(measure-sample)
                                              → ${var.sampleIndex}=0
                                              → Passed
  第1轮: bodyComplete? → false → resetBody() → advance() → sampleIndex=1
                                              → executeNode(measure-sample)
                                              → ${var.sampleIndex}=1
                                              → Passed
  第2轮: bodyComplete? → false → resetBody() → advance() → sampleIndex=2
                                              → executeNode(measure-sample)
                                              → ${var.sampleIndex}=2
                                              → Passed
  完成:  advance() → completed → outcome=Passed
  继续执行 after-loop""")

page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 10
# ═══════════════════════════════════════════════════════════════

add_heading('十、设计模式分析', 1)
add_table(
    ['设计模式', '代码位置', '解决的问题', '评价'],
    [
        ['策略模式', 'BarrierController: 6×4×5策略组合; ErrorPolicyEngine: 5种错误动作', '策略可从JSON配置切换，不需改调度器', '典型应用，通过枚举+switch实现，性能好但扩展需改枚举'],
        ['工厂方法', 'IDeviceSessionFactory, TransportDeviceSessionFactory', '设备session创建延迟到运行时', '标准抽象工厂用法'],
        ['适配器模式', 'TransportModuleAdapter: IModuleTransport适配为IModule', '进程/DLL对调度器看起来和内置模块一样', '整个模块隔离的核心设计'],
        ['注册表模式', 'ModuleRegistry, DeviceSessionManager的factory注册', '运行时动态注册模块/设备', '经典插件注册机制'],
        ['建造者模式', 'PlanBuilder: buildGroup/buildStep/buildLoopStep分层构建Plan', '编译阶段关注点分离', '构建过程清晰，各方法职责单一'],
        ['处理器链', 'NodeRunner中NoopHandler→WaitHandler→ActionHandler链', '不同节点类型由不同Handler处理', '轻量级责任链，可扩展新Handler'],
        ['管道模式', 'SequenceCompiler→PlanBuilder→ExecutionSession→ExecutionReport', 'JSON→模型→计划→执行→报告', '清晰的单向数据流'],
        ['不可变对象', 'ExecutionPlan使用shared_ptr<const ExecutionPlan>', '多UUT共享Plan无需加锁；缓存安全', '重要的并发安全设计'],
        ['DTO', 'ExecutionReport及子结构(纯数据，无行为)', 'UI只消费数据，不接触运行时对象', '三层解耦的关键'],
    ],
    [2.5, 5.5, 5.5, 5]
)

page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 11
# ═══════════════════════════════════════════════════════════════

add_heading('十一、扩展机制分析', 1)

add_heading('11.1 各扩展点成本评估', 2)
add_table(
    ['扩展点', '成本', '说明'],
    [
        ['新增业务模块', '极低', '实现IModule接口或遵循stdin/stdout JSON协议即可'],
        ['新增工站', '极低', '编写新的Station Config JSON文件，不改代码'],
        ['新增Transport类型', '低', '实现IModuleTransport接口即可'],
        ['新增设备类型', '中', '实现IDeviceSessionFactory + IDeviceSession + Station Config配置'],
        ['新增步骤类型', '高', '需改ExecNodeKind枚举 + SequenceCompiler + PlanBuilder + 新Handler'],
        ['新增Barrier策略', '高', '需改枚举和所有switch分支（建议策略接口改造）'],
        ['新增Loop类型 (While等)', '高', '需大量改动LoopController和PlanBuilder'],
    ],
    [4, 2.5, 12]
)

add_heading('11.2 开闭原则检查', 2)
add_table(
    ['扩展点', '对扩展开放', '对修改封闭', '评价'],
    [
        ['新增模块', '✓ IModule', '✓', '最成熟的扩展点'],
        ['新增Transport', '✓ IModuleTransport', '✓', '接口稳定'],
        ['新增设备驱动', '✓ IDeviceSessionFactory', '✓', '接口稳定'],
        ['新增工站', '✓ Station Config JSON', '✓', '纯配置'],
        ['新增Barrier策略', '✗ 需改枚举和switch', '✗', '可用策略模式改进'],
        ['新增错误动作', '✗ 需改ErrorAction枚举', '✗', '同上'],
        ['新增步骤类型', '✗ 需改多处', '△', '频率低，可接受'],
        ['新增变量类型', '△ RuntimeVariableResolver', '△', '需改两套实现'],
    ],
    [4, 5, 4.5, 5]
)

page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 12
# ═══════════════════════════════════════════════════════════════

add_heading('十二、异常处理设计', 1)

add_heading('12.1 异常分类与处理', 2)
add_table(
    ['异常类型', '处理机制', '代码位置'],
    [
        ['步骤失败 (Failed)', 'ErrorPolicyEngine 根据 onFail 策略处理', 'ErrorPolicyEngine.h'],
        ['模块错误 (Error)', 'ErrorPolicyEngine 根据 onError 策略处理', 'ErrorPolicyEngine.h'],
        ['执行超时 (Timeout)', 'ErrorPolicyEngine 根据 onTimeout 策略处理；Transport 层有独立 timeout', 'NodeRunner.cpp + 各Transport'],
        ['进程崩溃 (非零退出)', 'QProcessTransport 捕获 → TransportError', 'QProcessTransport.h'],
        ['DLL 超时', 'DllBridgeInvoker 线程超时 → DllExecuteTimeout', 'DllBridgeInvoker.h'],
        ['用户停止', 'requestStop → skipPendingNonAlwaysRun → activateAllCleanup', 'ExecutionSession.cpp:370'],
        ['资源获取超时', 'acquireTimeoutMs 已定义但当前未在执行循环中强制检查', 'ExecutionPlan.h:100'],
    ],
    [4, 8.5, 6]
)

add_heading('12.2 错误恢复策略', 2)
add_para('5 种错误动作（ExecutionPlan.h:87-93）：Continue（忽略继续） / StopUut（停止该UUT但执行Cleanup） / Retry（重试最多maxAttempts次） / RunCleanup（激活指定清理区） / Abort（立即中止整个Session）')

add_heading('12.3 资源释放保证', 2)
add_bullet('Step 级资源（ResourceManager）：在 executeNode() 的 finally 逻辑中保证释放（ExecutionGraphScheduler.cpp:360）')
add_bullet('Run/Station 级资源（DeviceSessionManager）：closeAll() 统一释放（DeviceSessionManager.h:88）')
add_bullet('模块进程：QProcessTransport 在析构时 kill；PersistentQProcessTransport 通过 shutdown() 显式关闭')

add_heading('12.4 卡死防护评估', 2)
add_table(
    ['防护机制', '实现状态', '评估'],
    [
        ['节点级别超时', 'TimeoutPolicy.timeoutMs 已定义', '当前NodeRunner未实际检查——需确认'],
        ['Transport 超时', 'QProcessTransport/DllBridge 有 timeoutMs', '已实现'],
        ['Barrier 超时', 'arrivalTimeoutMs + releaseTimeoutMs', '已实现'],
        ['资源申请超时', 'acquireTimeoutMs 已定义', '未在执行循环中检查过期'],
        ['停止兜底', 'requestStop(Abort) 可强制终止', '已实现'],
        ['看门狗/心跳', '无', '当前代码无法确认'],
    ],
    [4, 8, 6]
)
add_note('需要关注：当前执行循环是单线程 while 循环，如果某节点执行无限阻塞（如 Wait 节点设置很大的 ms），整个引擎会卡住。没有看门狗或心跳机制。')

page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 13
# ═══════════════════════════════════════════════════════════════

add_heading('十三、性能与并发分析', 1)

add_heading('13.1 线程模型', 2)
add_para('当前是单线程同步模型。整个调度循环在调用线程中执行（CLI 的主线程）。没有工作线程池、没有异步 Future/Promise、没有 Qt 信号槽用于执行。唯一的子线程使用在 DllBridgeInvoker 中用于 DLL 调用隔离。')
add_para('源码证据：ExecutionSession.cpp:260 的 while (progressed) 直接在调用线程；NodeRunner.cpp:129 的 QThread::msleep() 阻塞当前线程。')

add_heading('13.2 并发控制', 2)
add_para('所有并发控制都是逻辑层面的（资源互斥、Barrier 同步），而非线程层面的锁（无 QMutex、无 QWaitCondition）。')

add_heading('13.3 性能特征', 2)
add_table(
    ['特征', '评估'],
    [
        ['UUT 间并发', '伪并发：Round-robin 轮转，每次推进一个 UUT 的一个节点'],
        ['单 UUT 吞吐', '取决于模块执行耗时 + Wait 节点耗时'],
        ['多 UUT 吞吐', '模块执行快(<1ms)时几乎并行；模块耗时长时后续UUT等待'],
        ['内存占用', '轻量：每个 UUT 存储 NodeActivation 状态，Plan 共享'],
    ],
    [3.5, 15]
)

add_heading('13.4 潜在性能瓶颈', 2)
add_bullet('Wait 节点阻塞：所有 UUT 的主循环被 Wait 阻塞。如果 UUT-1 在 wait 10s，UUT-2 也无法执行其他步骤。')
add_bullet('无真正并发：多 UUT 不能同时执行耗时操作（如同时读 DMM）。')
add_bullet('资源等待是轮询：资源不可用时不是阻塞等待，而是每轮重试 tryAcquire——浪费 CPU。')
add_bullet('无优先级调度：UUT 之间没有优先级区分。')

page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 14
# ═══════════════════════════════════════════════════════════════

add_heading('十四、日志与可观测性', 1)

add_heading('14.1 当前机制', 2)
add_table(
    ['机制', '实现', '评价'],
    [
        ['traceId', 'TransportModuleAdapter 自动生成 moduleId:hex 格式', '跨进程追踪基础'],
        ['nodeId/attemptId', '所有运行时对象有唯一 ID', '可定位到具体步骤和尝试'],
        ['errorCode + errorMessage', 'ModuleResult / NodeResult / MeasurementResult 都有', '结构化错误信息'],
        ['stderr 捕获', 'QProcessTransport 可读取子进程 stderr', '基本诊断信息'],
        ['compileError.path', '编译错误带 JSON 路径（如 groups[0].steps[1].inputs）', '对排查JSON问题非常有用'],
    ],
    [3.5, 9, 6]
)

add_heading('14.2 当前不足', 2)
add_bullet('没有结构化日志系统（无 log level、无 logger 抽象）')
add_bullet('没有执行耗时统计（NodeResult 有 startedAt/finishedAt 但无耗时计算）')
add_bullet('没有指标收集（吞吐量、通过率、首次通过率等）')
add_bullet('没有分布式追踪（跨进程 trace 只有 traceId，无 span/parent 概念）')
add_bullet('CLI 的输出是 printf 风格，不是结构化的')

add_heading('14.3 现场排障评估', 2)
add_table(
    ['能力', '评估', '说明'],
    [
        ['JSON 配置错误定位', '好', '编译错误带精确 JSON 路径'],
        ['运行结果层级', '好', 'ExecutionReport 含完整 step/attempt/measurement 层级'],
        ['跨进程排查', '一般', '依赖 stderr 输出，没有统一日志格式'],
        ['性能分析', '不足', '没有时间线/火焰图式的性能数据'],
    ],
    [4, 2.5, 12]
)

page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 15
# ═══════════════════════════════════════════════════════════════

add_heading('十五、接口边界合理性评估', 1)

add_heading('15.1 逐项检查', 2)
add_table(
    ['检查项', '状态', '说明'],
    [
        ['UI 是否直接调用业务代码', '✓ 通过', '当前无UI；规划中UI只消费ExecutionReport'],
        ['UI 是否依赖业务对象', '✓ 通过', 'ExecutionReport 是纯 DTO'],
        ['业务是否直接更新 UI', '✓ 通过', '业务模块返回 ModuleResult，不知 UI 存在'],
        ['是否存在跨层调用', '✓ 通过', '引擎层通过抽象接口调用业务层'],
        ['是否存在双向依赖', '✓ 通过', '依赖方向单向：引擎 → IModule ← 业务实现'],
        ['是否违反依赖倒置原则', '✓ 通过', '引擎依赖抽象（IModule），业务实现抽象'],
        ['是否违反接口隔离原则', '△', 'IModuleRuntimeServices 接口较大但都是设备相关操作'],
        ['状态是否统一由引擎维护', '✓ 通过', '所有状态在 UutExecution/NodeActivation 中'],
        ['日志是否统一由引擎汇总', '△', '引擎汇总了errorCode/errorMessage，但无统一log channel'],
        ['结果是否统一由引擎管理', '✓ 通过', 'ExecutionReport 统一汇总'],
    ],
    [7.5, 3, 8]
)

add_heading('15.2 发现的问题', 2)

add_para('问题 1：没有 UI Facade 接口', bold=True)
add_para('当前 CLI 直接操作多个底层对象。未来 UI 接入时需要封装统一的 ITaskEngineFacade。')

add_para('问题 2：没有事件通知机制', bold=True)
add_para('UI 要获取状态变化需要轮询 state()/report()。应该增加 Observer/Callback 机制。')

add_para('问题 3：Pause/Resume 未实现', bold=True)
add_para('ExecutionState 中有 Paused 枚举，但 run() 中没有对应的暂停/恢复逻辑。')

add_para('问题 4：UutExecution 直接暴露给外部', bold=True)
add_para('uuts() 返回 QVector<UutExecution>&（可变引用），外部可以修改运行时状态，违反了封装原则。应只暴露 const 版本。源码：ExecutionSession.h:23。')

add_para('问题 5：DeviceSessionManager 直接暴露可变引用', bold=True)
add_para('devices() 返回 DeviceSessionManager&，允许外部绕过 Session 直接操作设备。源码：ExecutionSession.h:24。')

page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 16
# ═══════════════════════════════════════════════════════════════

add_heading('十六、优化建议', 1)

add_heading('16.1 建立统一 Facade 接口', 2)
add_para('建议封装 ITaskEngineFacade 作为 UI 的唯一入口：')
add_code("""class ITaskEngineFacade {
public:
    // 生命周期
    virtual LoadResult loadSequence(const QString& jsonPath) = 0;
    virtual LoadResult loadStation(const QString& stationPath) = 0;
    virtual void setUutCount(int count) = 0;

    // 控制
    virtual void start() = 0;
    virtual void pause() = 0;
    virtual void resume() = 0;
    virtual void stop(StopMode mode) = 0;

    // 查询
    virtual ExecutionState state() const = 0;
    virtual ExecutionReport report() const = 0;

    // 事件
    virtual void subscribe(IExecutionObserver* observer) = 0;
};""")
add_para('好处：UI 只依赖一个接口，内部复杂度被封装；切换引擎实现时 UI 不需要改动。')

add_heading('16.2 建立事件通知机制', 2)
add_code("""class IExecutionObserver {
public:
    virtual void onStateChanged(ExecutionState oldState, ExecutionState newState) = 0;
    virtual void onStepCompleted(const UutId& uutId, const StepReport& step) = 0;
    virtual void onProgressChanged(int completedSteps, int totalSteps) = 0;
};""")
add_para('好处：UI 不需要轮询，实时获得状态变化；适合未来 MES 的事件推送。')

add_heading('16.3 策略模式改造 Barrier/Error', 2)
add_para('将 BarrierController 中的枚举+switch 改为策略接口，新增策略不需要修改 BarrierController 代码。')

add_heading('16.4 日志体系建立', 2)
add_para('引入结构化日志 IExecutionLogger，每个 node/attempt 使用相同的 correlationId 串联所有日志。')

add_heading('16.5 其他优化建议', 2)
add_table(
    ['建议', '优先级', '说明'],
    [
        ['UutExecution 只暴露 const', '高', '封装运行时状态，防止外部误修改'],
        ['DeviceSessionManager 通过 Facade 控制', '高', '避免外部绕过 Session'],
        ['Pause/Resume 实现', '中', '当前枚举已定义但未实现'],
        ['Wait 节点不阻塞主循环', '中', '改用定时器或分片等待'],
        ['资源申请超时检查', '中', 'acquireTimeoutMs 已定义但未执行'],
        ['统一变量替换接口', '低', 'Configuration 和 Runtime 合并抽象'],
        ['多线程并发执行', '低', '需仔细设计，对现有模型冲击大'],
    ],
    [6, 2.5, 10]
)

page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 17 — APPENDIX
# ═══════════════════════════════════════════════════════════════

add_heading('十七、附录', 1)

add_heading('17.1 核心类速查', 2)
add_table(
    ['类名', '文件', '职责'],
    [
        ['ExecutionPlan', 'ExecutionPlan.h', '不可变执行计划'],
        ['UutExecution', 'RuntimeTypes.h', '单 UUT 运行时状态'],
        ['NodeActivation', 'RuntimeTypes.h', '节点激活状态'],
        ['NodeAttempt', 'RuntimeTypes.h', '单次执行尝试'],
        ['ExecutionSession', 'ExecutionSession.h', '多 UUT 执行会话'],
        ['ExecutionGraphScheduler', 'ExecutionGraphScheduler.h', '节点调度器'],
        ['NodeRunner', 'NodeRunner.h', '节点分发执行器'],
        ['ResourceManager', 'ResourceManager.h', '资源仲裁'],
        ['BarrierController', 'BarrierController.h', '批次同步'],
        ['LoopController', 'LoopController.h', '循环控制'],
        ['ErrorPolicyEngine', 'ErrorPolicyEngine.h', '错误策略决策'],
        ['SequenceCompiler', 'SequenceCompiler.h', 'JSON 编译器'],
        ['PlanBuilder', 'PlanBuilder.h', '执行计划构建器'],
        ['ModuleRegistry', 'ModuleRuntime.h', '模块注册表'],
        ['TransportModuleAdapter', 'ModuleRuntime.h', 'Transport→Module 适配器'],
        ['DeviceSessionManager', 'DeviceSessionManager.h', '设备连接管理'],
        ['StationRuntime', 'StationRuntime.h', '工站运行上下文'],
        ['VariableResolver', 'VariableResolver.h', '配置期变量替换'],
        ['RuntimeVariableResolver', 'RuntimeVariableResolver.h', '运行期变量替换'],
        ['ExecutionReport', 'ExecutionReport.h', '只读运行报告'],
    ],
    [5, 7.5, 6]
)

add_heading('17.2 核心接口速查', 2)
add_table(
    ['接口', '文件', '方法数', '用途'],
    [
        ['IModule', 'ModuleRuntime.h', '2', '业务模块接入'],
        ['IModuleTransport', 'ModuleRuntime.h', '1', 'Transport 抽象'],
        ['INodeHandler', 'NodeRunner.h', '2', '节点处理器'],
        ['IModuleRuntimeServices', 'ModuleRuntime.h', '4', '设备服务注入'],
        ['IDeviceSession', 'DeviceSessionManager.h', '7', '设备会话'],
        ['IDeviceSessionFactory', 'DeviceSessionManager.h', '2', '设备会话工厂'],
        ['IDeviceCommandSession', 'ModuleRuntime.h', '1', '设备命令代理'],
        ['ISessionPersistence', 'SessionSnapshot.h', '2', '会话持久化（占位）'],
    ],
    [5, 5, 2, 6.5]
)

add_heading('17.3 枚举速查', 2)
add_table(
    ['枚举', '值数量', '文件'],
    [
        ['ExecutionState', '9', 'RuntimeTypes.h'],
        ['ActivationState', '11', 'RuntimeTypes.h'],
        ['AttemptState', '4', 'RuntimeTypes.h'],
        ['FrameState', '4', 'RuntimeTypes.h'],
        ['ExecNodeKind', '6', 'ExecutionPlan.h'],
        ['EdgeKind', '4', 'ExecutionPlan.h'],
        ['EdgeTrigger', '9', 'ExecutionPlan.h'],
        ['NodeOutcome', '7', 'ExecutionPlan.h'],
        ['ErrorAction', '5', 'ExecutionPlan.h'],
        ['ResourceMode', '5', 'ExecutionPlan.h'],
        ['CleanupReason', '8', 'ExecutionPlan.h'],
        ['ModuleOutcome', '4', 'ModuleRuntime.h'],
        ['ModuleTransportStatus', '3', 'ModuleRuntime.h'],
        ['MeasurementStatus', '5', 'MeasurementTypes.h'],
        ['BarrierArrivalPolicy', '6', 'BarrierController.h'],
        ['BarrierReleasePolicy', '4', 'BarrierController.h'],
        ['BarrierFailurePolicy', '5', 'BarrierController.h'],
        ['BarrierTimeoutPolicy', '5', 'BarrierController.h'],
        ['DeviceSessionLifetime', '3', 'DeviceSessionManager.h'],
        ['DeviceConnectionState', '4', 'DeviceSessionManager.h'],
    ],
    [5.5, 2.5, 5.5]
)

add_heading('17.4 目录结构总览', 2)
add_code("""PicoATE/
├── docs/          (15 份设计/规范/总结文档)
├── examples/      (16 个示例 JSON/Python/Manifest)
├── src/
│   ├── core/      (27 头文件 + 27 实现 = 核心引擎, 静态库)
│   ├── cli/       (1 文件 = 命令行入口, 517行)
│   ├── mockhost/  (1 文件 = 测试 Echo 进程, 72行)
│   ├── nativehost/(1 文件 = DLL 隔离加载器, 207行)
│   ├── fakeinstrumenthost/ (1 文件 = 假仪器长驻进程, 320行)
│   ├── testdllmodule/ (1 文件 = C ABI 测试 DLL, 80行)
│   └── canexamplemodule/ (1 文件 = 模拟 CAN DLL, 247行)
├── tests/         (1 测试文件 3408行 = 34个单元测试 + 12个ctest)
└── out/           (VS2022 构建输出)""")

add_heading('17.5 测试覆盖概览', 2)
add_table(
    ['测试方向', '覆盖内容', '用例数'],
    [
        ['ResourceManager', '资源互斥、等待队列序列化', '1'],
        ['BarrierController', 'Barrier 成员到达、跨 UUT 释放、DropFailed member', '2'],
        ['PlanCache', 'shared_ptr<const ExecutionPlan> 引用语义', '1'],
        ['Scheduler', 'Retry、Cleanup、Barrier、Stop 语义', '4'],
        ['SequenceDef / PlanBuilder', 'Setup/Main/Cleanup 建模、disabled、custom group、loop region', '8'],
        ['SequenceCompiler', 'JSON编译、类型校验、未知字段警告、loop错误、module bindings', '8'],
        ['Example File E2E', '全部9个example JSON端到端运行', '9'],
        ['ExecutionReport', '结果DTO覆盖retry和失败cleanup', '4'],
        ['Module Runtime', '注册执行、缺失错误、outputs/measurements映射', '2'],
        ['Module Transport', 'Adapter成功/超时/错误映射', '3'],
        ['QProcessTransport', 'MockHost正常/超时/非零退出', '3'],
        ['Module Bindings', '变量替换错误报告', '1'],
        ['DllBridgeInvoker', 'DLL成功/错误码/超时', '3'],
        ['NativeHost', 'DLL echo + manifest + timeout kill', '3'],
        ['CAN DLL', '模拟CAN解码pass + limit fail', '2'],
        ['VariableResolver', '配置期+运行期变量替换', '4'],
        ['DeviceSessionManager', '连接复用、缺driver、connect失败', '2'],
        ['StationConfig', '解析、变量替换、disabled device、错误lifetime', '2'],
        ['StationRuntime + CLI', '工站加载 + --station 示例', '1'],
        ['Persistent Host', '长驻进程状态保持、health/reconnect/shutdown', '2'],
        ['DMM/CAN Adapter', 'adapter spike全链路', '1'],
    ],
    [3.5, 12, 2]
)

add_para('')
add_para('── 文档结束 ──', bold=True, size=11)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('本架构设计文档基于 PicoATE 源码（截至 2026-06-26）完整分析得出。\n标注"当前代码无法确认"的项目代表源码中没有对应实现或设计文档中没有相关规划。')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

# ── Save ────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'PicoATE_Architecture_Design_Document.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
